import docx
import nltk
from nltk.translate.meteor_score import meteor_score
from tabulate import tabulate

# Download necessary NLTK data
nltk.download('punkt')
nltk.download('wordnet')

# Function to read text from a DOCX file
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Tokenization function
def tokenize(text):
    return nltk.word_tokenize(text.lower())

# Paths to the DOCX files containing the original text and reference summaries
original_text_path = "/content/drive/MyDrive/Colab Notebooks/datasets/a_letter_to_god.docx"
reference_summaries_path = '/content/drive/MyDrive/Colab Notebooks/datasets/ref_summ.docx'

# Read and tokenize the original text
original_text = read_docx(original_text_path)
original_text_tokens = tokenize(original_text)

# Read and tokenize the reference summaries from the DOCX file
reference_summaries_doc = docx.Document(reference_summaries_path)
reference_summaries = [tokenize(para.text) for para in reference_summaries_doc.paragraphs if para.text.strip()]

# Add additional reference summaries directly
additional_references = [
    "In the heart of all who lived in that solitary house in the middle of the valley, there was a single hope: help from God . The following Sunday, at daybreak, he began to write a letter which he himself would carry The following Sunday Lencho came a bit earlier than usual to ask if there was a letter for him It was the postman himself who handed the letter to him . Lencho showed not the slightest surprise on seeing the money",
    "lencho was an ox of a man, working like an animal in the fields, but still he knew how to write . he wrote a letter to God, asking for money from his employees, he himself gave part of his salary, and several friends of his were obliged to give something ‘for an act of charity."
]

# Tokenize additional reference summaries and extend the existing list
reference_summaries.extend([tokenize(ref) for ref in additional_references])

# Generated summaries
generated_summaries = {
    'google/flan-t5-base': 'Lencho was an ox of a man, working like an animal in the fields, but still he knew how to write. The following Sunday, at daybreak, he began to write a letter to God.',
    'google/t5-base': 'lencho was an ox of a man, working like an animal in the fields, but still he knew how to write . he wrote a letter to God, asking for money from his employees, he himself gave part of his salary, and several friends of his were obliged to give something ‘for an act of charity’ . when he opened the letter, it was evident that to answer it he needed something more than goodwill,',
    'facebook/bart-large-cnn': 'Lencho, who knew his fields intimately, saw the sky towards the north-east. A strong wind began to blow and along with the rain very large hailstones began to fall. For an hour the hail rained on the house, the garden, the hillside, the cornfield, on the whole valley.',
    'sshleifer/distilbart-cnn-12-6': 'In the heart of all who lived in that solitary house in the middle of the valley, there was a single hope: help from God . The following Sunday, at daybreak, he began to write a letter which he himself would carry The following Sunday Lencho came a bit earlier than usual to ask if there was a letter for him It was the postman himself who handed the letter to him . Lencho showed not the slightest surprise on seeing the money; such was',
}

# List to hold the results
results = []

# Calculate METEOR score for each summary
for model_name, summary in generated_summaries.items():
    # Tokenize the summary
    summary_tokens = tokenize(summary)
    # Calculate the METEOR score using multiple references
    score = meteor_score(reference_summaries, summary_tokens)
    results.append([model_name, f'{score:.4f}'])

# Print the results in a table
headers = ["Model", "METEOR Score"]
print(tabulate(results, headers, tablefmt="grid"))
